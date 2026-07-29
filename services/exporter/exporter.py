import io
import requests
from docx import Document
from docx.shared import Inches, Pt
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY

def export_to_txt(title, content_html_or_text):
    """Compiles title and story content into a clean text file stream."""
    mem_file = io.BytesIO()

    # Simple plain-text compilation layout
    txt_content = f"{title.upper()}\n\n{'='*len(title)}\n\n{content_html_or_text}"
    mem_file.write(txt_content.encode('utf-8'))
    mem_file.seek(0)
    return mem_file

def export_to_docx(title, plain_content, thumbnail_url=None):
    """Renders a dynamic document layout template in memory."""
    doc = Document()
    
    #Document Title
    title_p = doc.add_paragraph()
    title_run = title_p.add_run(title)
    title_run.font.size = Pt(24)
    title_run.bold = True
    title_p.alignment = 1
    
    #Render Thumbnail Image if it exists
    if thumbnail_url:
        try:
            response = requests.get(thumbnail_url, timeout=5)
            if response.status_code == 200:
                img_stream = io.BytesIO(response.content)
                doc.add_picture(img_stream, width=Inches(5))
                doc.add_paragraph().alignment = 1
        except Exception as e:
            print(f"Skipping DOCX image insert due to error: {e}")
            
    #Content Body Template
    doc.add_paragraph(plain_content)
    
    mem_file = io.BytesIO()
    doc.save(mem_file)
    mem_file.seek(0)
    return mem_file

def export_to_pdf(title, plain_content, thumbnail_url=None):
    """Generates a dynamic PDF document on-the-fly with styled layout elements."""
    mem_file = io.BytesIO()
    doc = SimpleDocTemplate(mem_file, pagesize=letter, rightMargin=54, leftMargin=54, topMargin=54, bottomMargin=54)
    story_elements = []
    
    styles = getSampleStyleSheet()
    
    # Custom Layout Typography Styles
    title_style = ParagraphStyle(
        'CustomTitle', parent=styles['Heading1'], fontSize=26, leading=30, alignment=TA_CENTER, spaceAfter=20
    )
    body_style = ParagraphStyle(
        'CustomBody', parent=styles['BodyText'], fontSize=11, leading=16, alignment=TA_JUSTIFY, spaceAfter=12
    )

    #Title Element
    story_elements.append(Paragraph(title, title_style))
    story_elements.append(Spacer(1, 15))
    
    #Dynamic Image Element Rendering
    if thumbnail_url:
        try:
            response = requests.get(thumbnail_url, timeout=5)
            if response.status_code == 200:
                img_stream = io.BytesIO(response.content)

                # Aspect ratio configuration matching 854x480 resolution bounds cleanly
                story_elements.append(Image(img_stream, width=400, height=225))
                story_elements.append(Spacer(1, 20))
        except Exception as e:
            print(f"Skipping PDF image insert due to error: {e}")
            
    #Narrative Paragraph Elements
    paragraphs = plain_content.split('\n\n')
    for p in paragraphs:
        if p.strip():
            story_elements.append(Paragraph(p.strip(), body_style))
            
    doc.build(story_elements)
    mem_file.seek(0)
    return mem_file
